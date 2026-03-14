import matplotlib 
matplotlib.use("Agg")
from flask import Flask, render_template, request, jsonify, send_from_directory, Response
from flask_wtf.csrf import CSRFProtect
import fitz
import os
import tempfile
import re
import subprocess
import html
import uuid
import time
import threading
import traceback
from werkzeug.utils import secure_filename
import hashlib
import shutil
import zipfile
import json
import warnings

# Suppress version-related warnings
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", message=".*LibreSSL.*")

# ReportLab Imports
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Frame, PageTemplate, Preformatted, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.pdfgen import canvas
from io import BytesIO

# python-docx Imports
from docx import Document
from docx.shared import Pt, Inches

# Additional Imports
import matplotlib.pyplot as plt
import requests

# Local Imports
from context_generator import ContextGenerator
from highlight_extractor import HighlightExtractor

app = Flask(__name__, static_folder='static', template_folder='templates')
temp_dir = tempfile.mkdtemp()
app.config['UPLOAD_FOLDER'] = temp_dir
app.config['SECRET_KEY'] = 'luminar-secret-key-2025'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100 MB limit

plt.switch_backend('agg')
csrf = CSRFProtect(app)

file_timestamps = {}
CLEANUP_INTERVAL = 3600 # 1 hour

def track_file_access(filename):
    if filename: file_timestamps[filename] = time.time()

def cleanup_aged_files():
    current_time = time.time()
    files_to_remove = [k for k, v in list(file_timestamps.items()) if current_time - v > CLEANUP_INTERVAL]
    for filename in files_to_remove:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        try:
            if os.path.exists(filepath):
                if os.path.isdir(filepath): shutil.rmtree(filepath)
                else: os.remove(filepath)
            del file_timestamps[filename]
        except Exception as e: print(f"Cleanup Error: {e}")

def cleanup_task():
    while True:
        time.sleep(600)
        cleanup_aged_files()

threading.Thread(target=cleanup_task, daemon=True).start()

# --- UTILITY & CORE LOGIC FUNCTIONS ---
def to_roman(n):
    if not isinstance(n, int) or n <= 0: return str(n)
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    res = ''
    for i in range(len(val)):
        while n >= val[i]: res += syb[i]; n -= val[i]
    return res

def to_alpha(n, uppercase=True):
    if not isinstance(n, int) or n <= 0: return str(n)
    result = ""; start = 65 if uppercase else 97
    while n > 0:
        n, r = divmod(n - 1, 26)
        result = chr(start + r) + result
    return result

def get_doc_stats(filepath):
    try:
        doc = fitz.open(filepath); stats = {'pages': len(doc), 'size': os.path.getsize(filepath)}; doc.close()
        return stats
    except Exception: return {'pages': 0, 'size': 0}

conversion_status = {}

@app.route('/check_conversion_status/<filename>')
def check_conversion_status(filename):
    return jsonify(conversion_status.get(filename, {'status': 'unknown'}))

def _convert_ipynb_to_pdf_async(src_path):
    server_key = os.path.basename(src_path); pdf_output_path = os.path.splitext(src_path)[0] + '.pdf'
    try:
        conversion_status[server_key] = {'status': 'processing', 'progress': 5}
        proc = subprocess.Popen(['jupyter', 'nbconvert', '--to', 'webpdf', '--allow-chromium-download', src_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        progress = 10
        while proc.poll() is None:
            time.sleep(2); progress = min(90, progress + 5)
            conversion_status[server_key]['progress'] = progress
        if proc.returncode == 0 and os.path.exists(pdf_output_path):
            conversion_status[server_key] = {'status': 'done', 'pdf_path': pdf_output_path, 'progress': 100, 'pdf_basename': os.path.basename(pdf_output_path)}
            track_file_access(os.path.basename(pdf_output_path))
        else:
            _, err = proc.communicate(timeout=5)
            conversion_status[server_key] = {'status': 'failed', 'error': err.decode('utf-8', errors='ignore') or 'nbconvert failed.'}
    except Exception as e: conversion_status[server_key] = {'status': 'failed', 'error': str(e)}

def get_pdf_for_serverfile(server_filename, input_filepath, wait_seconds=15):
    if server_filename.lower().endswith('.pdf'): return input_filepath
    start_wait = time.time()
    while time.time() - start_wait < wait_seconds:
        status = conversion_status.get(server_filename, {})
        if status.get('status') == 'done': return status.get('pdf_path')
        if status.get('status') == 'failed' or status.get('status') == 'unknown': break
        time.sleep(1)
    return conversion_status.get(server_filename, {}).get('pdf_path')

def extract_highlights(pdf_path):
    doc = fitz.open(pdf_path); extractor = HighlightExtractor(doc); extracted_data = extractor.extract_highlights()
    highlights = [(item.get('type', 'point'), item.get('text', '')) for item in extracted_data]
    doc.close(); return highlights

def create_modern_pdf(highlights, output_path):
    # SimpleDocTemplate is more robust than BaseDocTemplate for general usage
    doc = SimpleDocTemplate(output_path, pagesize=letter, leftMargin=0.75*inch, rightMargin=0.75*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='ModernHeading', fontName='Helvetica-Bold', fontSize=16, textColor=colors.HexColor("#2c3e50"), alignment=TA_CENTER, spaceAfter=15, leading=20))
    styles.add(ParagraphStyle(name='ModernBody', fontName='Helvetica', fontSize=12, textColor=colors.black, spaceAfter=10, leading=16))
    styles.add(ParagraphStyle(name='ModernCode', fontName='Courier', fontSize=11, textColor=colors.HexColor("#2c3e50"), backColor=colors.HexColor("#f8f9fa"), borderColor=colors.HexColor("#dee2e6"), borderWidth=1, borderPadding=8, leading=14, spaceAfter=12))

    def add_page_branding(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#3b82f6"))
        canvas.setFont('Helvetica-Bold', 10)
        canvas.drawCentredString(letter[0]/2, letter[1] - 0.5*inch, "HEPHAESTUS STUDY NOTES")
        canvas.setFillColor(colors.gray)
        canvas.setFont('Helvetica', 8)
        canvas.drawCentredString(letter[0]/2, 0.5*inch, f"Page {doc.page}")
        canvas.restoreState()

    story = [Paragraph("Generated Study Signals", styles['ModernHeading']), Spacer(1, 20)]
    for item_type, text in highlights:
        if not text or not text.strip(): continue
        if item_type == 'heading': story.append(Paragraph(text, styles['ModernHeading']))
        elif item_type == 'code': story.append(Preformatted(text, styles['ModernCode']))
        elif item_type == 'math':
            try:
                fig = plt.figure(figsize=(6, 1), facecolor='white'); fig.text(0.5, 0.5, f'${text}$', ha='center', va='center', fontsize=20, color='black')
                img_path = os.path.join(temp_dir, f'math_{uuid.uuid4()}.png'); plt.savefig(img_path, transparent=True, bbox_inches='tight', pad_inches=0.1); plt.close(fig)
                story.append(Image(img_path, width=4*inch, height=0.5*inch)); os.remove(img_path)
            except Exception: story.append(Paragraph(text, styles['ModernBody']))
        else: story.append(Paragraph(f'<bullet color="#3b82f6">•</bullet> {html.escape(text)}', styles['ModernBody']))
        story.append(Spacer(1, 8))
    
    doc.build(story, onFirstPage=add_page_branding, onLaterPages=add_page_branding)

def create_docx_from_highlights(highlights, output_path):
    doc = Document(); title = doc.add_paragraph(); run = title.add_run('Study Notes'); run.bold = True; run.font.size = Pt(18)
    for item_type, text in highlights:
        if not text or not text.strip(): continue
        if item_type == 'heading': p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(14)
        elif item_type == 'code': p = doc.add_paragraph(); r = p.add_run(text); r.font.name = 'Courier New'; r.font.size = Pt(10)
        else: doc.add_paragraph('• ' + text)
    doc.save(output_path)

def create_docx_from_pdf(pdf_path: str, docx_output_path: str) -> None:
    try:
        pdf = fitz.open(pdf_path); doc = Document()
        for page in pdf:
            pix = page.get_pixmap(dpi=144); img_path = os.path.join(temp_dir, f"p_{uuid.uuid4()}.png"); pix.save(img_path); doc.add_picture(img_path, width=Inches(6.5)); os.remove(img_path)
        doc.save(docx_output_path)
    except Exception as e: print(f"DOCX conversion failed: {e}")

def add_header_footer_to_pdf(input_pdf_path, output_filepath, headers, footers, start_page_num, page_num_placement, page_num_format, overlap_resolution, margin_size, chapter_num, page_num_enabled, hf_enabled):
    input_doc = fitz.open(input_pdf_path); output_doc = fitz.open()
    header_y_pos = letter[1] - 0.5*inch; footer_y_pos = 0.5*inch; hf_x_margin = 0.5*inch; content_margin = 0.1*inch
    total_pages = len(input_doc); last_page_num = start_page_num + total_pages - 1
    page_num_area, page_num_pos = page_num_placement.split('-')
    for i, page in enumerate(input_doc):
        temp_headers = headers.copy(); temp_footers = footers.copy()
        new_page = output_doc.new_page(width=page.rect.width, height=page.rect.height)
        packet = BytesIO(); can = canvas.Canvas(packet, pagesize=(page.rect.width, page.rect.height)); can.setFont('Helvetica', 9)
        curr = start_page_num + i; p_str = ""
        if page_num_enabled:
            fmt_map = {'roman_lower': to_roman(curr).lower(), 'roman_upper': to_roman(curr), 'alpha_lower': to_alpha(curr, False), 'alpha_upper': to_alpha(curr), 'dash_x_dash': f"- {curr} -", 'page_x': f"Page {curr}", 'page_x_of_n': f"Page {curr} of {last_page_num}"}
            p_str = fmt_map.get(page_num_format, str(curr))
            target = temp_headers if page_num_area == 'header' else temp_footers
            if target.get(page_num_pos):
                if overlap_resolution == 'before': target[page_num_pos] = f"{p_str} {target[page_num_pos]}"
                else: target[page_num_pos] = f"{target[page_num_pos]} {p_str}"
                p_str = ""
        if hf_enabled:
            can.drawString(hf_x_margin, header_y_pos, temp_headers.get('left', '')); can.drawCentredString(page.rect.width / 2, header_y_pos, temp_headers.get('center', '')); can.drawRightString(page.rect.width - hf_x_margin, header_y_pos, temp_headers.get('right', ''))
            can.drawString(hf_x_margin, footer_y_pos, temp_footers.get('left', '')); can.drawCentredString(page.rect.width / 2, footer_y_pos, temp_footers.get('center', '')); can.drawRightString(page.rect.width - hf_x_margin, footer_y_pos, temp_footers.get('right', ''))
        if p_str:
            y = header_y_pos if page_num_area == 'header' else footer_y_pos
            if page_num_pos == 'left': can.drawString(hf_x_margin, y, p_str)
            elif page_num_pos == 'right': can.drawRightString(page.rect.width - hf_x_margin, y, p_str)
            else: can.drawCentredString(page.rect.width / 2, y, p_str)
        can.save(); packet.seek(0); overlay_doc = fitz.open("pdf", packet.read()); new_page.show_pdf_page(new_page.rect, overlay_doc, 0)
        content_rect = fitz.Rect(content_margin, content_margin, page.rect.width - content_margin, page.rect.height - content_margin); new_page.show_pdf_page(content_rect, input_doc, i)
    output_doc.save(output_filepath); output_doc.close(); input_doc.close()

# --- ROUTES ---
@app.route('/')
def home(): return render_template('index.html')

@app.route('/temp/<path:filename>')
def serve_temp_file(filename):
    track_file_access(filename)
    as_attachment = 'download' in request.args
    # Fix: serve Markdown as text/plain so it's viewable in browser
    if not as_attachment and filename.lower().endswith('.md'):
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename, mimetype='text/plain')
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=as_attachment)

@csrf.exempt
@app.route('/upload_and_analyze', methods=['POST'])
def upload_and_analyze():
    if 'file' not in request.files: return jsonify({'error': 'No file'}), 400
    f = request.files['file']; original = secure_filename(f.filename); content = f.read()
    filename = f"{hashlib.sha1(content).hexdigest()[:12]}_{original}"; path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(path):
        with open(path, 'wb') as out: out.write(content)
    track_file_access(filename)
    if filename.lower().endswith('.ipynb'): threading.Thread(target=_convert_ipynb_to_pdf_async, args=(path,)).start()
    return jsonify({'serverFilename': filename, 'initialStats': get_doc_stats(path)})

@csrf.exempt
@app.route('/extract_highlights', methods=['POST'])
def extract_highlights_route():
    try:
        s_name = request.form.get('serverFilename'); path = os.path.join(app.config['UPLOAD_FOLDER'], s_name); pdf_path = get_pdf_for_serverfile(s_name, path)
        if not pdf_path or not os.path.exists(pdf_path): return jsonify({'error': 'PDF not ready'}), 404
        highlights = extract_highlights(pdf_path)
        if not highlights: return jsonify({'error': 'No highlights'}), 400
        base = os.path.splitext(os.path.basename(pdf_path))[0]; pdf_out = os.path.join(app.config['UPLOAD_FOLDER'], f"{base}_notes.pdf"); docx_out = os.path.join(app.config['UPLOAD_FOLDER'], f"{base}_notes.docx")
        create_modern_pdf(highlights, pdf_out); create_docx_from_highlights(highlights, docx_out); track_file_access(os.path.basename(pdf_out)); track_file_access(os.path.basename(docx_out))
        return jsonify({'previewUrl': f'/temp/{os.path.basename(pdf_out)}', 'docxUrl': f'/temp/{os.path.basename(docx_out)}', 'finalStats': get_doc_stats(pdf_out)})
    except Exception as e: traceback.print_exc(); return jsonify({'error': str(e)}), 500

@csrf.exempt
@app.route('/add_header_footer', methods=['POST'])
def add_header_footer_route():
    try:
        d = request.form; s_name = d.get('serverFilename'); path = os.path.join(app.config['UPLOAD_FOLDER'], s_name); pdf_path = get_pdf_for_serverfile(s_name, path)
        if not pdf_path or not os.path.exists(pdf_path): return jsonify({'error': 'PDF not ready'}), 404
        headers = {'left': d.get('headerLeft',''), 'center': d.get('headerCenter',''), 'right': d.get('headerRight','')}
        footers = {'left': d.get('footerLeft',''), 'center': d.get('footerCenter',''), 'right': d.get('footerRight','')}
        output = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4()}_final.pdf")
        add_header_footer_to_pdf(pdf_path, output, headers, footers, int(d.get('startPageNum',1)), d.get('pageNumPlacement','footer-center'), d.get('pageNumFormat','numeric'), 'after', 'normal', '1', d.get('isPageNumEnabled')=='true', d.get('isHfEnabled')=='true')
        docx = output.replace('.pdf', '.docx'); create_docx_from_pdf(output, docx); track_file_access(os.path.basename(output)); track_file_access(os.path.basename(docx))
        return jsonify({'previewUrl': f'/temp/{os.path.basename(output)}', 'docxUrl': f'/temp/{os.path.basename(docx)}', 'finalStats': get_doc_stats(output)})
    except Exception as e: traceback.print_exc(); return jsonify({'error': str(e)}), 500

@csrf.exempt
@app.route('/generate_context', methods=['POST'])
def generate_context_route():
    try:
        f = request.files['file']; fmt = request.form.get('outputFormat', 'html')
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4()}_{secure_filename(f.filename)}"); f.save(zip_path); extract_path = os.path.join(app.config['UPLOAD_FOLDER'], f"ex_{uuid.uuid4()}")
        with zipfile.ZipFile(zip_path, 'r') as z: z.extractall(extract_path)
        items = os.listdir(extract_path); a_path = os.path.join(extract_path, items[0]) if items and os.path.isdir(os.path.join(extract_path, items[0])) else extract_path
        generator = ContextGenerator(); content = generator.generate_context_from_folder(a_path, fmt)
        ext = 'html' if fmt == 'html' else 'md'; out_name = f"ctx_{uuid.uuid4()}.{ext}"; out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_name)
        with open(out_path, 'w', encoding='utf-8') as out_f: out_f.write(content)
        track_file_access(out_name); d_name = f"{os.path.splitext(f.filename)[0]}_context.{ext}"
        return jsonify({'previewUrl': f'/temp/{out_name}', 'downloadUrl': f'/temp/{out_name}?download=true&filename={d_name}', 'filename': d_name})
    except Exception as e: traceback.print_exc(); return jsonify({'error': str(e)}), 500
    finally: shutil.rmtree(extract_path, ignore_errors=True); os.remove(zip_path) if os.path.exists(zip_path) else None

if __name__ == '__main__': app.run(host="0.0.0.0", port=5001, debug=True)
